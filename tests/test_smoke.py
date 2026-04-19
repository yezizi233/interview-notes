from __future__ import annotations

import argparse
import tempfile
import unittest
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / "interview-notes" / "scripts"
import sys

sys.path.insert(0, str(SCRIPTS_DIR))

from interview_notes import run_pipeline  # noqa: E402


class SmokeTest(unittest.TestCase):
    def test_transcript_pipeline_generates_two_docx_files(self) -> None:
        transcript_path = REPO_ROOT / "examples" / "sample_transcript.txt"
        with tempfile.TemporaryDirectory() as tmpdir:
            args = argparse.Namespace(
                transcript=str(transcript_path),
                audio=None,
                template=None,
                title=None,
                slug="smoke-test",
                outdir=tmpdir,
                output_language="en",
                provider="none",
                model=None,
                network_mode="offline",
                whisper_model="base",
            )
            report = run_pipeline(args)

            clear_path = Path(report["clear_docx"])
            trace_path = Path(report["traceability_docx"])

            self.assertTrue(clear_path.exists(), "clear docx was not created")
            self.assertTrue(trace_path.exists(), "traceability docx was not created")
            self.assertGreater(report["trace_entry_count"], 0)

            clear_doc = Document(str(clear_path))
            trace_doc = Document(str(trace_path))
            self.assertTrue(any(p.text.strip() for p in clear_doc.paragraphs))
            self.assertTrue(trace_doc.tables, "traceability docx should contain at least one table")


if __name__ == "__main__":
    unittest.main()
