from __future__ import annotations

import argparse
import json
import math
import os
import re
import tempfile
import textwrap
import wave
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Optional

import requests
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


STOPWORDS_EN = {
    "the",
    "and",
    "for",
    "that",
    "with",
    "from",
    "this",
    "have",
    "will",
    "your",
    "about",
    "there",
    "their",
    "they",
    "were",
    "into",
    "which",
    "would",
    "could",
    "should",
    "also",
    "after",
    "before",
    "being",
    "when",
    "where",
    "what",
    "while",
    "than",
    "then",
    "because",
    "through",
    "across",
    "under",
    "over",
    "between",
    "interview",
    "notes",
    "topic",
    "part",
    "section",
}

STOPWORDS_ZH = {
    "我们",
    "他们",
    "这个",
    "那个",
    "以及",
    "因为",
    "所以",
    "如果",
    "但是",
    "进行",
    "关于",
    "一个",
    "一些",
    "已经",
    "可以",
    "需要",
    "就是",
    "不是",
    "没有",
    "还有",
    "比较",
    "目前",
    "时候",
    "情况",
}

AMBIGUOUS_MARKERS = [
    "???",
    "[inaudible]",
    "(inaudible)",
    "[unclear]",
    "(unclear)",
    "听不清",
    "不清楚",
    "待确认",
]

SPEAKER_TIMESTAMP_PATTERNS = [
    re.compile(
        r"^\s*\[(?P<ts>\d{1,2}:\d{2}(?::\d{2})?)\]\s*(?:(?P<speaker>[^:：\-\]]{1,60})\s*[:：\-]\s*)?(?P<text>.+?)\s*$"
    ),
    re.compile(
        r"^\s*(?P<ts>\d{1,2}:\d{2}(?::\d{2})?)\s+(?:(?P<speaker>[^:：\-]{1,60})\s*[:：\-]\s*)?(?P<text>.+?)\s*$"
    ),
    re.compile(
        r"^\s*(?P<speaker>[^:：\[]+?)\s*\[(?P<ts>\d{1,2}:\d{2}(?::\d{2})?)\]\s*[:：\-]?\s*(?P<text>.+?)\s*$"
    ),
    re.compile(
        r"^\s*(?P<speaker>[A-Za-z0-9\u4e00-\u9fff _/&().-]{1,40})\s*[:：]\s*(?P<text>.+?)\s*$"
    ),
]

CHINESE_NUMERALS = "一二三四五六七八九十"


@dataclass
class Segment:
    segment_id: str
    index: int
    text: str
    timestamp: Optional[str] = None
    start_seconds: Optional[float] = None
    end_seconds: Optional[float] = None
    speaker: Optional[str] = None
    source_medium: str = "transcript"
    original_text: Optional[str] = None


@dataclass
class TemplateProfile:
    path: Optional[Path] = None
    output_language: Optional[str] = None
    title_style: Optional[str] = None
    section_style: Optional[str] = None
    topic_style: Optional[str] = None
    bullet_style: Optional[str] = None
    subbullet_style: Optional[str] = None
    table_style: str = "Table Grid"
    section_prefix_mode: str = "zh"


@dataclass
class NameCheck:
    raw: str
    normalized: str
    status: str
    authority: str = ""
    note: str = ""


@dataclass
class TraceEntry:
    entry_id: str
    clean_location: str
    source_medium: str
    speaker: str
    timestamp: str
    raw_excerpt: str
    final_wording: str
    transformation_method: str
    name_verification: str
    uncertainty: str


def detect_network() -> bool:
    try:
        response = requests.get(
            "https://clinicaltrials.gov/api/query/study_fields?expr=NCT00000102&fields=NCTId&min_rnk=1&max_rnk=1&fmt=json",
            timeout=5,
        )
        return response.ok
    except requests.RequestException:
        return False


def parse_timecode(value: str) -> Optional[float]:
    if not value:
        return None
    parts = value.split(":")
    try:
        parts_i = [int(part) for part in parts]
    except ValueError:
        return None
    if len(parts_i) == 2:
        minutes, seconds = parts_i
        return (minutes * 60) + seconds
    if len(parts_i) == 3:
        hours, minutes, seconds = parts_i
        return (hours * 3600) + (minutes * 60) + seconds
    return None


def seconds_to_timecode(seconds: Optional[float]) -> str:
    if seconds is None:
        return "timestamp unavailable"
    total = max(0, int(round(seconds)))
    hours, rem = divmod(total, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def slugify(value: str) -> str:
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE).strip().lower()
    value = re.sub(r"[\s_-]+", "-", value)
    return value[:80] or "interview-notes"


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    chunks = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def read_text_input(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    return path.read_text(encoding="utf-8")


def derive_language_from_text(text: str) -> str:
    zh_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    en_count = len(re.findall(r"[A-Za-z]", text))
    if zh_count >= en_count:
        return "zh"
    return "en"


def parse_transcript(text: str, source_medium: str = "transcript") -> list[Segment]:
    segments: list[Segment] = []
    current: Optional[dict[str, Any]] = None
    lines = text.splitlines()
    idx = 0

    def flush_current() -> None:
        nonlocal idx, current
        if not current:
            return
        payload = current.copy()
        payload["text"] = " ".join(part.strip() for part in payload["parts"] if part.strip()).strip()
        if payload["text"]:
            idx += 1
            segments.append(
                Segment(
                    segment_id=f"S{idx:04d}",
                    index=idx,
                    text=payload["text"],
                    original_text=payload["text"],
                    timestamp=payload.get("timestamp"),
                    start_seconds=payload.get("start_seconds"),
                    speaker=payload.get("speaker"),
                    source_medium=source_medium,
                )
            )
        current = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_current()
            continue

        matched = False
        for pattern in SPEAKER_TIMESTAMP_PATTERNS:
            match = pattern.match(line)
            if not match:
                continue
            groups = match.groupdict()
            timestamp = groups.get("ts")
            speaker = groups.get("speaker")
            text_part = groups.get("text", "").strip()
            flush_current()
            current = {
                "timestamp": timestamp,
                "start_seconds": parse_timecode(timestamp) if timestamp else None,
                "speaker": speaker.strip() if speaker else None,
                "parts": [text_part],
            }
            matched = True
            break

        if matched:
            continue

        if current is None:
            current = {
                "timestamp": None,
                "start_seconds": None,
                "speaker": None,
                "parts": [line],
            }
        else:
            current["parts"].append(line)

    flush_current()

    for i, segment in enumerate(segments):
        next_start = segments[i + 1].start_seconds if i + 1 < len(segments) else None
        if segment.start_seconds is not None and next_start is not None and next_start > segment.start_seconds:
            segment.end_seconds = next_start
        elif segment.start_seconds is not None:
            segment.end_seconds = segment.start_seconds + 20
    return segments


def analyze_template(path: Optional[Path]) -> TemplateProfile:
    profile = TemplateProfile(path=path)
    if path is None:
        return profile

    doc = Document(str(path))
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    profile.output_language = derive_language_from_text("\n".join(p.text for p in nonempty[:30])) if nonempty else None
    style_names = {style.name for style in doc.styles}
    if "Title" in style_names:
        profile.title_style = "Title"
    if "Heading 1" in style_names:
        profile.section_style = "Heading 1"
    if "Heading 2" in style_names:
        profile.topic_style = "Heading 2"
    if "List Bullet" in style_names:
        profile.bullet_style = "List Bullet"
    if "List Bullet 2" in style_names:
        profile.subbullet_style = "List Bullet 2"

    for paragraph in nonempty:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else None
        if profile.title_style is None:
            profile.title_style = style_name
        if re.match(rf"^[{CHINESE_NUMERALS}]+、", text) or re.match(r"^\d+\.", text):
            profile.section_style = style_name or profile.section_style
            if re.match(rf"^[{CHINESE_NUMERALS}]+、", text):
                profile.section_prefix_mode = "zh"
            else:
                profile.section_prefix_mode = "en"
            break

    for paragraph in nonempty:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else None
        if text.startswith("【") and text.endswith("】"):
            profile.topic_style = style_name or profile.topic_style
            break

    for table in doc.tables:
        style_name = table.style.name if table.style is not None else None
        if style_name:
            profile.table_style = style_name
            break

    return profile


def load_template_document(profile: TemplateProfile) -> Document:
    if profile.path and profile.path.suffix.lower() == ".docx":
        doc = Document(str(profile.path))
        body = doc._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        return doc
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    return doc


def extract_candidate_terms(text: str) -> list[str]:
    terms = set(re.findall(r"\bNCT\d{8}\b", text))
    for match in re.findall(r"\b[A-Z][A-Za-z0-9]+(?:\s+[A-Z][A-Za-z0-9&()./-]+){0,3}\b", text):
        if len(match) > 2:
            terms.add(match.strip())
    return sorted(terms)


def verify_terms(terms: Iterable[str], network_available: bool) -> tuple[dict[str, NameCheck], list[str]]:
    notices: list[str] = []
    results: dict[str, NameCheck] = {}
    if not network_available:
        notices.append("Network unavailable: online proper-name verification is disabled.")
        for term in terms:
            results[term] = NameCheck(raw=term, normalized=term, status="unverified offline")
        return results, notices

    for term in terms:
        if re.fullmatch(r"NCT\d{8}", term):
            try:
                response = requests.get(
                    "https://clinicaltrials.gov/api/query/study_fields",
                    params={
                        "expr": term,
                        "fields": "NCTId,BriefTitle,OfficialTitle",
                        "min_rnk": 1,
                        "max_rnk": 1,
                        "fmt": "json",
                    },
                    timeout=10,
                )
                response.raise_for_status()
                payload = response.json()
                studies = payload.get("StudyFieldsResponse", {}).get("StudyFields", [])
                if studies:
                    study = studies[0]
                    normalized = study.get("OfficialTitle", []) or study.get("BriefTitle", []) or [term]
                    normalized_text = normalized[0]
                    results[term] = NameCheck(
                        raw=term,
                        normalized=normalized_text,
                        status="verified online",
                        authority="ClinicalTrials.gov",
                    )
                    continue
            except requests.RequestException as exc:
                notices.append(f"ClinicalTrials.gov lookup failed for {term}: {exc}")

        results[term] = NameCheck(raw=term, normalized=term, status="unverified", note="No authoritative adapter configured.")

    if any(value.status == "unverified" for value in results.values()):
        notices.append(
            "Some company/product/institution names remain unverified. Current automatic verification supports NCT trial IDs; other terms are preserved and flagged as unverified."
        )
    return results, notices


def apply_name_normalization(text: str, checks: dict[str, NameCheck]) -> str:
    normalized = text
    for raw, info in sorted(checks.items(), key=lambda item: len(item[0]), reverse=True):
        if info.status == "verified online" and raw != info.normalized:
            normalized = normalized.replace(raw, info.normalized)
    return normalized


def choose_provider(provider: str) -> str:
    if provider != "auto":
        return provider
    if os.getenv("OPENAI_API_KEY"):
        return "openai"
    if os.getenv("ANTHROPIC_API_KEY"):
        return "anthropic"
    return "none"


def call_openai(system_prompt: str, user_prompt: str, model: str) -> dict[str, Any]:
    api_key = os.environ["OPENAI_API_KEY"]
    base = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/")
    response = requests.post(
        f"{base}/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": model,
            "temperature": 0.1,
            "response_format": {"type": "json_object"},
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        },
        timeout=120,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    return json.loads(content)


def call_anthropic(system_prompt: str, user_prompt: str, model: str) -> dict[str, Any]:
    api_key = os.environ["ANTHROPIC_API_KEY"]
    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": model,
            "max_tokens": 4000,
            "temperature": 0.1,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        },
        timeout=120,
    )
    response.raise_for_status()
    content_blocks = response.json()["content"]
    text = "".join(block["text"] for block in content_blocks if block.get("type") == "text")
    return json.loads(text)


def build_llm_prompts(
    title: str,
    output_language: str,
    template_profile: TemplateProfile,
    notices: list[str],
    segments: list[Segment],
    checks: dict[str, NameCheck],
) -> tuple[str, str]:
    system_prompt = textwrap.dedent(
        """
        You generate faithful interview notes.

        Hard rules:
        - Do not fabricate.
        - Do not omit supported points.
        - Keep uncertainty when present.
        - Organize into interview notes, not meeting action items.
        - Output valid JSON only.
        - Every bullet or table row must cite one or more source segment IDs.
        - Prefer topic-based organization.
        - If output_language is zh, write Chinese. If output_language is en, write English.
        - Use verified normalized names when provided. Leave unverified names unchanged.
        - Use tables only for numeric-heavy comparisons.
        """
    ).strip()

    template_summary = {
        "output_language": output_language,
        "section_prefix_mode": template_profile.section_prefix_mode,
        "has_template": bool(template_profile.path),
    }
    checks_summary = [
        {
            "raw": check.raw,
            "normalized": check.normalized,
            "status": check.status,
            "authority": check.authority,
        }
        for check in checks.values()
    ]
    user_payload = {
        "title": title,
        "template_profile": template_summary,
        "notices": notices,
        "name_checks": checks_summary,
        "segments": [
            {
                "segment_id": segment.segment_id,
                "timestamp": segment.timestamp or "timestamp unavailable",
                "speaker": segment.speaker or "unknown",
                "text": segment.text,
            }
            for segment in segments
        ],
        "schema": {
            "title": "string",
            "sections": [
                {
                    "title": "string",
                    "topics": [
                        {
                            "title": "string",
                            "bullets": [
                                {
                                    "text": "string",
                                    "source_segment_ids": ["S0001"],
                                    "subpoints": [
                                        {
                                            "text": "string",
                                            "source_segment_ids": ["S0002"],
                                        }
                                    ],
                                }
                            ],
                            "tables": [
                                {
                                    "title": "string",
                                    "columns": ["c1", "c2"],
                                    "rows": [
                                        {
                                            "cells": ["v1", "v2"],
                                            "source_segment_ids": ["S0003"],
                                        }
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ],
            "notes": ["optional pipeline notices to show in traceability file"],
        },
    }
    user_prompt = json.dumps(user_payload, ensure_ascii=False)
    return system_prompt, user_prompt


def try_llm_generation(
    provider: str,
    model: Optional[str],
    title: str,
    output_language: str,
    template_profile: TemplateProfile,
    notices: list[str],
    segments: list[Segment],
    checks: dict[str, NameCheck],
) -> tuple[Optional[dict[str, Any]], list[str]]:
    provider_name = choose_provider(provider)
    if provider_name == "none":
        return None, ["No LLM provider configured. Falling back to local evidence-preserving formatter."]

    system_prompt, user_prompt = build_llm_prompts(title, output_language, template_profile, notices, segments, checks)

    try:
        if provider_name == "openai":
            payload = call_openai(system_prompt, user_prompt, model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini"))
        else:
            payload = call_anthropic(system_prompt, user_prompt, model or os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-latest"))
        return payload, [f"Structured note generation used provider: {provider_name}."]
    except Exception as exc:  # noqa: BLE001
        return None, [f"LLM generation failed and the pipeline fell back to local formatting: {exc}"]


def keyword_topic_title(text: str, index: int, language: str) -> str:
    if language == "zh":
        words = re.findall(r"[\u4e00-\u9fff]{2,}", text)
        scores: dict[str, int] = {}
        for word in words:
            if word in STOPWORDS_ZH:
                continue
            scores[word] = scores.get(word, 0) + 1
        top = sorted(scores, key=scores.get, reverse=True)[:2]
        if top:
            return f"【{' / '.join(top)}】"
        return f"【主题 {index}】"

    words = re.findall(r"[A-Za-z][A-Za-z0-9/-]{2,}", text)
    scores = {}
    for word in words:
        lower = word.lower()
        if lower in STOPWORDS_EN:
            continue
        scores[word] = scores.get(word, 0) + 1
    top = sorted(scores, key=scores.get, reverse=True)[:3]
    if top:
        return f"[{' / '.join(top)}]"
    return f"[Topic {index}]"


def fallback_structure(title: str, output_language: str, segments: list[Segment]) -> dict[str, Any]:
    if not segments:
        return {"title": title, "sections": [], "notes": ["No transcript content was parsed."]}

    chunk_size = max(4, math.ceil(len(segments) / 4))
    section_title = "一、访谈内容整理" if output_language == "zh" else "1. Interview Content"
    topics = []
    for topic_idx, start in enumerate(range(0, len(segments), chunk_size), 1):
        chunk = segments[start : start + chunk_size]
        chunk_text = " ".join(segment.text for segment in chunk)
        bullets = []
        for segment in chunk:
            bullets.append(
                {
                    "text": segment.text,
                    "source_segment_ids": [segment.segment_id],
                    "subpoints": [],
                }
            )
        topics.append(
            {
                "title": keyword_topic_title(chunk_text, topic_idx, output_language),
                "bullets": bullets,
                "tables": [],
            }
        )

    note = (
        "Local fallback formatter was used. The clean file preserves source-grounded statements but may be less synthesized than an LLM-assisted run."
    )
    return {"title": title, "sections": [{"title": section_title, "topics": topics}], "notes": [note]}


def attach_audio_verification(
    segments: list[Segment],
    audio_path: Optional[Path],
    whisper_model: str,
) -> tuple[list[Segment], list[str]]:
    if audio_path is None:
        return segments, []

    notices: list[str] = []
    try:
        import av
        from faster_whisper import WhisperModel
    except Exception as exc:  # noqa: BLE001
        return segments, [f"Audio verification unavailable: {exc}"]

    def extract_clip(input_path: Path, start_sec: float, end_sec: float) -> Path:
        temp_fd, temp_path = tempfile.mkstemp(suffix=".wav")
        os.close(temp_fd)
        out_path = Path(temp_path)
        with av.open(str(input_path)) as container, wave.open(str(out_path), "wb") as wav_out:
            stream = container.streams.audio[0]
            resampler = av.audio.resampler.AudioResampler(format="s16", layout="mono", rate=16000)
            wav_out.setnchannels(1)
            wav_out.setsampwidth(2)
            wav_out.setframerate(16000)
            for frame in container.decode(stream):
                frame_time = float(frame.time or 0.0)
                frame_end = frame_time + (frame.samples / float(frame.sample_rate or 16000))
                if frame_end < start_sec:
                    continue
                if frame_time > end_sec:
                    break
                resampled = resampler.resample(frame)
                frames = resampled if isinstance(resampled, list) else [resampled]
                for item in frames:
                    if item is None:
                        continue
                    wav_out.writeframes(item.to_ndarray().tobytes())
        return out_path

    ambiguous_segments = [
        segment
        for segment in segments
        if any(marker.lower() in segment.text.lower() for marker in AMBIGUOUS_MARKERS)
    ]
    if not ambiguous_segments:
        return segments, []

    model = WhisperModel(whisper_model, device="cpu", compute_type="int8")
    for segment in ambiguous_segments:
        if segment.start_seconds is None or segment.end_seconds is None:
            notices.append(
                f"Skipped narrow audio verification for {segment.segment_id}: transcript segment has no timestamp bounds."
            )
            continue
        clip_start = max(0.0, segment.start_seconds - 2)
        clip_end = max(clip_start + 3, segment.end_seconds + 2)
        clip_path = extract_clip(audio_path, clip_start, clip_end)
        try:
            parts, _ = model.transcribe(str(clip_path), beam_size=5)
            corrected = " ".join(part.text.strip() for part in parts if part.text.strip()).strip()
            if corrected:
                segment.original_text = segment.text
                segment.text = corrected
                segment.source_medium = "transcript + recording"
                notices.append(f"Audio-corrected ambiguous segment {segment.segment_id}.")
        finally:
            clip_path.unlink(missing_ok=True)

    return segments, notices


def transcribe_audio(audio_path: Path, whisper_model: str) -> tuple[list[Segment], list[str]]:
    try:
        from faster_whisper import WhisperModel
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Audio transcription unavailable: {exc}") from exc

    model = WhisperModel(whisper_model, device="cpu", compute_type="int8")
    parts, _ = model.transcribe(str(audio_path), beam_size=5, vad_filter=True)
    segments: list[Segment] = []
    for idx, part in enumerate(parts, 1):
        segments.append(
            Segment(
                segment_id=f"S{idx:04d}",
                index=idx,
                text=part.text.strip(),
                original_text=part.text.strip(),
                timestamp=seconds_to_timecode(part.start),
                start_seconds=float(part.start),
                end_seconds=float(part.end),
                speaker=None,
                source_medium="recording",
            )
        )
    return segments, [f"Full audio transcription used faster_whisper model: {whisper_model}."]


def detect_name_status(text: str, checks: dict[str, NameCheck]) -> str:
    matched = [info for raw, info in checks.items() if raw in text or info.normalized in text]
    if not matched:
        return "unverified"
    priorities = {"verified online": 2, "unverified": 1, "unverified offline": 0}
    return sorted(matched, key=lambda item: priorities.get(item.status, -1), reverse=True)[0].status


def build_trace_entries(
    payload: dict[str, Any],
    segments_by_id: dict[str, Segment],
    checks: dict[str, NameCheck],
) -> list[TraceEntry]:
    entries: list[TraceEntry] = []
    counter = 1

    def make_entry(location: str, wording: str, source_ids: list[str], method: str) -> None:
        nonlocal counter
        if not source_ids:
            source_ids = []
        refs = [segments_by_id[source_id] for source_id in source_ids if source_id in segments_by_id]
        raw_excerpt = " | ".join(ref.text for ref in refs)[:1000] if refs else ""
        source_medium = " + ".join(sorted({ref.source_medium for ref in refs})) if refs else "transcript"
        speaker_values = [ref.speaker for ref in refs if ref.speaker]
        speaker = speaker_values[0] if len(set(speaker_values)) == 1 and speaker_values else "not reliably labeled"
        timestamps = [ref.timestamp for ref in refs if ref.timestamp]
        timestamp = " - ".join([timestamps[0], timestamps[-1]]) if len(timestamps) > 1 else (timestamps[0] if timestamps else "timestamp unavailable")
        uncertainty = "none"
        if any(ref.original_text and ref.original_text != ref.text for ref in refs):
            method = "audio-corrected"
        if timestamp == "timestamp unavailable":
            uncertainty = "timestamp unavailable"
        if speaker == "not reliably labeled":
            uncertainty = "speaker uncertain" if uncertainty == "none" else uncertainty
        entries.append(
            TraceEntry(
                entry_id=f"T-{counter:04d}",
                clean_location=location,
                source_medium=source_medium,
                speaker=speaker,
                timestamp=timestamp,
                raw_excerpt=raw_excerpt,
                final_wording=wording,
                transformation_method=method,
                name_verification=detect_name_status(wording, checks),
                uncertainty=uncertainty,
            )
        )
        counter += 1

    for s_idx, section in enumerate(payload.get("sections", []), 1):
        for t_idx, topic in enumerate(section.get("topics", []), 1):
            for b_idx, bullet in enumerate(topic.get("bullets", []), 1):
                location = f"{section['title']} > {topic['title']} > bullet {b_idx}"
                make_entry(location, bullet["text"], bullet.get("source_segment_ids", []), "compressed rewrite")
                for sb_idx, subpoint in enumerate(bullet.get("subpoints", []), 1):
                    location = f"{section['title']} > {topic['title']} > bullet {b_idx}.{sb_idx}"
                    make_entry(location, subpoint["text"], subpoint.get("source_segment_ids", []), "compressed rewrite")
            for table_idx, table in enumerate(topic.get("tables", []), 1):
                for row_idx, row in enumerate(table.get("rows", []), 1):
                    location = f"{section['title']} > {topic['title']} > table {table_idx} row {row_idx}"
                    make_entry(location, " | ".join(row.get("cells", [])), row.get("source_segment_ids", []), "table-normalized")
    return entries


def add_paragraph_with_style(doc: Document, text: str, style_name: Optional[str], fallback: Optional[str] = None):
    paragraph = doc.add_paragraph()
    chosen_style = None
    for candidate in [style_name, fallback]:
        if candidate:
            try:
                paragraph.style = candidate
                chosen_style = candidate
                break
            except KeyError:
                continue
    if not chosen_style and fallback == "Title":
        run = paragraph.add_run()
        run.bold = True
        run.font.size = Pt(16)
    paragraph.add_run(text)
    return paragraph


def render_clean_docx(
    out_path: Path,
    payload: dict[str, Any],
    profile: TemplateProfile,
) -> None:
    doc = load_template_document(profile)
    add_paragraph_with_style(doc, payload.get("title", "Interview Notes"), profile.title_style, "Title")

    for section in payload.get("sections", []):
        add_paragraph_with_style(doc, section["title"], profile.section_style, "Heading 1")
        for topic in section.get("topics", []):
            add_paragraph_with_style(doc, topic["title"], profile.topic_style, "Heading 2")
            for bullet in topic.get("bullets", []):
                add_paragraph_with_style(doc, bullet["text"], profile.bullet_style, "List Bullet")
                for subpoint in bullet.get("subpoints", []):
                    add_paragraph_with_style(doc, subpoint["text"], profile.subbullet_style, "List Bullet 2")
            for table in topic.get("tables", []):
                if table.get("title"):
                    add_paragraph_with_style(doc, table["title"], profile.topic_style, "Heading 3")
                columns = table.get("columns", [])
                rows = table.get("rows", [])
                if not columns:
                    continue
                tbl = doc.add_table(rows=1, cols=len(columns))
                try:
                    tbl.style = profile.table_style
                except KeyError:
                    tbl.style = "Table Grid"
                for idx, column in enumerate(columns):
                    tbl.rows[0].cells[idx].text = column
                for row in rows:
                    cells = tbl.add_row().cells
                    for idx, value in enumerate(row.get("cells", [])):
                        cells[idx].text = str(value)

    doc.save(str(out_path))


def render_traceability_docx(
    out_path: Path,
    title: str,
    notices: list[str],
    entries: list[TraceEntry],
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    add_paragraph_with_style(doc, f"{title} - Traceability", None, "Title")
    if notices:
        add_paragraph_with_style(doc, "Pipeline Notices", None, "Heading 1")
        for notice in notices:
            add_paragraph_with_style(doc, notice, None, "List Bullet")

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "Entry ID",
        "Clean File Location",
        "Source Medium",
        "Speaker",
        "Timestamp",
        "Raw Excerpt",
        "Final Wording",
        "Transformation Method",
        "Name Verification",
        "Uncertainty",
    ]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for entry in entries:
        row = table.add_row().cells
        values = [
            entry.entry_id,
            entry.clean_location,
            entry.source_medium,
            entry.speaker,
            entry.timestamp,
            entry.raw_excerpt,
            entry.final_wording,
            entry.transformation_method,
            entry.name_verification,
            entry.uncertainty,
        ]
        for idx, value in enumerate(values):
            row[idx].text = value

    doc.save(str(out_path))


def determine_output_language(
    requested: str,
    transcript_text: str,
    template_profile: TemplateProfile,
) -> str:
    if requested in {"zh", "en"}:
        return requested
    if template_profile.output_language in {"zh", "en"}:
        return template_profile.output_language
    return derive_language_from_text(transcript_text)


def section_prefix(index: int, language: str, mode: str) -> str:
    if language == "zh" or mode == "zh":
        numerals = list(CHINESE_NUMERALS)
        if 0 < index <= len(numerals):
            return f"{numerals[index - 1]}、"
        return f"{index}、"
    return f"{index}."


def normalize_section_titles(payload: dict[str, Any], output_language: str, profile: TemplateProfile) -> dict[str, Any]:
    for idx, section in enumerate(payload.get("sections", []), 1):
        prefix = section_prefix(idx, output_language, profile.section_prefix_mode)
        title = section.get("title", "").strip()
        if not re.match(rf"^[{CHINESE_NUMERALS}\d]+[、.]", title):
            section["title"] = f"{prefix} {title}".strip()
    return payload


def build_title(args: argparse.Namespace, transcript_text: str) -> str:
    if args.title:
        return args.title
    if args.transcript:
        return Path(args.transcript).stem
    if args.audio:
        return Path(args.audio).stem
    first_line = next((line.strip() for line in transcript_text.splitlines() if line.strip()), "Interview Notes")
    return first_line[:80]


def run_pipeline(args: argparse.Namespace) -> dict[str, Any]:
    notices: list[str] = []
    template_profile = analyze_template(Path(args.template) if args.template else None)

    transcript_text = ""
    segments: list[Segment] = []
    if args.transcript:
        transcript_text = read_text_input(Path(args.transcript))
        segments = parse_transcript(transcript_text, source_medium="transcript")
    elif args.audio:
        segments, audio_notices = transcribe_audio(Path(args.audio), args.whisper_model)
        transcript_text = "\n".join(segment.text for segment in segments)
        notices.extend(audio_notices)
    else:
        raise ValueError("At least one of --transcript or --audio is required.")

    if args.audio and args.transcript:
        segments, audio_notices = attach_audio_verification(segments, Path(args.audio), args.whisper_model)
        notices.extend(audio_notices)
        transcript_text = "\n".join(segment.text for segment in segments)

    network_available = detect_network() if args.network_mode == "auto" else args.network_mode == "online"
    if not network_available and args.network_mode != "offline":
        notices.append(
            "Offline or verification-disabled mode: authoritative online proper-name verification is unavailable."
        )

    output_language = determine_output_language(args.output_language, transcript_text, template_profile)
    if output_language != derive_language_from_text(transcript_text) and choose_provider(args.provider) == "none":
        notices.append(
            "Requested output language differs from the detected source language, but no LLM provider is configured. The fallback formatter preserves source wording and may not fully translate."
        )

    candidate_terms = extract_candidate_terms(transcript_text)
    name_checks, name_notices = verify_terms(candidate_terms, network_available)
    notices.extend(name_notices)

    title = build_title(args, transcript_text)
    title = apply_name_normalization(title, name_checks)

    llm_payload, llm_notices = try_llm_generation(
        args.provider,
        args.model,
        title,
        output_language,
        template_profile,
        notices,
        segments,
        name_checks,
    )
    notices.extend(llm_notices)
    payload = llm_payload if llm_payload is not None else fallback_structure(title, output_language, segments)
    payload = normalize_section_titles(payload, output_language, template_profile)

    segments_by_id = {segment.segment_id: segment for segment in segments}
    trace_entries = build_trace_entries(payload, segments_by_id, name_checks)

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    slug = args.slug or slugify(title)
    clear_path = outdir / f"{slug}_clear.docx"
    trace_path = outdir / f"{slug}_traceability.docx"

    render_clean_docx(clear_path, payload, template_profile)
    render_traceability_docx(trace_path, payload.get("title", title), notices + payload.get("notes", []), trace_entries)

    report = {
        "title": payload.get("title", title),
        "output_language": output_language,
        "network_available": network_available,
        "clear_docx": str(clear_path),
        "traceability_docx": str(trace_path),
        "segment_count": len(segments),
        "trace_entry_count": len(trace_entries),
        "notices": notices + payload.get("notes", []),
    }
    return report
